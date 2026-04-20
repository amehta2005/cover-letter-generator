"""
exporter.py — Fill the DOCX template and export to DOCX + PDF.

Template placeholder system
────────────────────────────
The template.docx should contain these tokens anywhere in its paragraphs:

    {{DATE}}              — letter date
    {{SALUTATION}}        — "Dear Hiring Manager," / "Dear Jane Smith,"
    {{OPENING}}           — opening paragraph
    {{BODY1}}             — first body paragraph
    {{BODY2}}             — second body paragraph
    {{CLOSING}}           — closing paragraph
    {{NAME}}              — sender's full name

If the template does NOT contain these placeholders, the exporter falls back
to a "body-block" strategy: it identifies paragraphs that look like body text
(not headers, not addresses, not the signature line) and replaces them in
sequence with the generated paragraphs.

PDF generation
──────────────
Tries LibreOffice headless first (cross-platform, free), then docx2pdf.
If neither is available, returns None for the PDF path but still saves the DOCX.
"""

import copy
import os
import re
import shutil
import subprocess
import tempfile
from datetime import date
from typing import Optional

from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt

import config


# ── Placeholder tokens ────────────────────────────────────────────────────────

PLACEHOLDERS = ["{{DATE}}", "{{SALUTATION}}", "{{OPENING}}", "{{BODY1}}", "{{BODY2}}", "{{CLOSING}}", "{{NAME}}"]

SECTION_KEYS = {
    "{{DATE}}": "date",
    "{{SALUTATION}}": "salutation",
    "{{OPENING}}": "opening",
    "{{BODY1}}": "body1",
    "{{BODY2}}": "body2",
    "{{CLOSING}}": "closing",
    "{{NAME}}": "name",
}


# ── Low-level DOCX paragraph helpers ─────────────────────────────────────────

def _para_full_text(para) -> str:
    """Join all run texts in a paragraph (handles run-splitting)."""
    return "".join(r.text for r in para.runs)


def _replace_para_text(para, new_text: str) -> None:
    """
    Replace the entire text of a paragraph while preserving the first run's
    character formatting (font, size, bold, italic, color).

    Strategy:
      1. Save formatting from the first non-empty run.
      2. Set run[0].text = new_text
      3. Clear all other runs (set text = "")
    """
    if not para.runs:
        para.add_run(new_text)
        return

    # Find first run with actual content to borrow its format
    fmt_run = None
    for r in para.runs:
        if r.text.strip():
            fmt_run = r
            break
    if fmt_run is None:
        fmt_run = para.runs[0]

    # Clear all runs
    for r in para.runs:
        r.text = ""

    # Write new text into the format run
    fmt_run.text = new_text


def _insert_paragraph_after(para, text: str, style_name: str = None) -> None:
    """
    Insert a new paragraph immediately after `para` in the document,
    copying the style of `para` unless style_name is provided.
    """
    new_para = copy.deepcopy(para._p)
    para._p.addnext(new_para)
    # Now `new_para` is an lxml element; wrap it
    from docx.oxml import OxmlElement
    new_p = para._p.getnext()

    # Clear runs in the cloned paragraph and set new text
    from docx.text.paragraph import Paragraph
    wrapped = Paragraph(new_p, para._parent)
    _replace_para_text(wrapped, text)


# ── Template detection ────────────────────────────────────────────────────────

def _has_placeholders(doc: Document) -> bool:
    for para in doc.paragraphs:
        if any(ph in _para_full_text(para) for ph in PLACEHOLDERS):
            return True
    return False


def _find_body_range(doc: Document) -> tuple[int, int]:
    """
    Find the exact paragraph indices of the body block by anchoring on
    structural landmarks rather than counting from the edges:
      - Body START: first paragraph after the salutation line (Dear ...)
      - Body END:   paragraph before the closing line (Sincerely / Best ...)

    This handles templates of any length without fragile index offsets.
    """
    paras = doc.paragraphs
    salutation_idx = None
    closing_idx    = None

    for i, para in enumerate(paras):
        text = para.text.strip()
        if not text:
            continue
        if salutation_idx is None and re.match(r'^Dear\b', text, re.I):
            salutation_idx = i
        # Detect closing AFTER we've found a salutation
        if salutation_idx is not None and i > salutation_idx:
            if re.match(
                r'^(Sincerely|Best|Regards|Warm\s+regards|Kind\s+regards|'
                r'Thank\s+you|With\s+appreciation|Respectfully)',
                text, re.I
            ):
                closing_idx = i
                break

    start = (salutation_idx + 1) if salutation_idx is not None else 5
    end   = closing_idx           if closing_idx   is not None else len(paras) - 2
    return start, end


# ── Main fill function ────────────────────────────────────────────────────────

def fill_template(
    template_path: str,
    sections: dict[str, str],
    output_path: str,
) -> str:
    """
    Fill `template_path` with generated `sections` and save to `output_path`.
    Returns the output_path.

    sections keys: date, salutation, opening, body1, body2, closing, name
    """
    doc = Document(template_path)

    if _has_placeholders(doc):
        _fill_by_placeholders(doc, sections)
    else:
        _fill_by_body_block(doc, sections)

    doc.save(output_path)
    return output_path


def _fill_by_placeholders(doc: Document, sections: dict[str, str]) -> None:
    """Replace placeholder tokens in the document with generated text."""
    for para in doc.paragraphs:
        full_text = _para_full_text(para)
        for placeholder, key in SECTION_KEYS.items():
            if placeholder in full_text and key in sections:
                new_text = full_text.replace(placeholder, sections[key])
                _replace_para_text(para, new_text)
                break  # each paragraph holds at most one placeholder

    # Also scan tables (some templates put content in tables)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    full_text = _para_full_text(para)
                    for placeholder, key in SECTION_KEYS.items():
                        if placeholder in full_text and key in sections:
                            new_text = full_text.replace(placeholder, sections[key])
                            _replace_para_text(para, new_text)
                            break


def _fill_by_body_block(doc: Document, sections: dict[str, str]) -> None:
    """
    Fallback: replace every paragraph between the salutation and the closing
    with the generated content, deleting any extras.

    Uses structural anchoring (Dear ... / Sincerely ...) rather than
    index offsets, so it works on templates of any length.
    """
    paras     = list(doc.paragraphs)  # snapshot — we'll mutate the doc below
    start, end = _find_body_range(doc)
    body_paras = paras[start:end]

    generated = [g for g in [
        sections.get("opening", ""),
        sections.get("body1", ""),
        sections.get("body2", ""),
        sections.get("body3", ""),
        sections.get("closing", ""),
    ] if g.strip()]

    for slot_idx, para in enumerate(body_paras):
        if slot_idx < len(generated):
            _replace_para_text(para, generated[slot_idx])
        else:
            # Template has more slots than we need — delete the extras
            p_elem = para._element
            parent = p_elem.getparent()
            if parent is not None:
                parent.remove(p_elem)

    # If we have more generated paragraphs than template slots, insert the extras
    # after the last filled slot, copying that slot's paragraph style.
    if len(generated) > len(body_paras) and body_paras:
        last_para = body_paras[min(len(body_paras), len(generated)) - 1]
        for extra_text in generated[len(body_paras):]:
            # Clone the last paragraph element and insert it after
            import copy as _copy
            new_p = _copy.deepcopy(last_para._element)
            last_para._element.addnext(new_p)
            from docx.text.paragraph import Paragraph as _Para
            new_para = _Para(new_p, last_para._parent)
            _replace_para_text(new_para, extra_text)
            last_para = new_para  # chain insertions in order

    # Replace date line if present in the header area
    for para in paras[:start]:
        text = para.text.strip()
        if re.match(r'^\d{1,2}[/\-\.]\d{1,2}[/\-\.]\d{2,4}$', text) or \
           re.match(r'^(January|February|March|April|May|June|July|'
                    r'August|September|October|November|December)', text, re.I):
            _replace_para_text(para, sections.get("date", text))
            break


# ── PDF export ────────────────────────────────────────────────────────────────

def _find_libreoffice() -> Optional[str]:
    for path in config.LIBREOFFICE_PATHS:
        if os.path.isfile(path) and os.access(path, os.X_OK):
            return path
        # Also try shutil.which for entries that are just command names
        found = shutil.which(path)
        if found:
            return found
    return None


def export_pdf_libreoffice(docx_path: str, output_dir: str) -> Optional[str]:
    """Convert DOCX to PDF using LibreOffice headless."""
    lo = _find_libreoffice()
    if not lo:
        return None
    try:
        result = subprocess.run(
            [lo, "--headless", "--convert-to", "pdf", "--outdir", output_dir, docx_path],
            capture_output=True, text=True, timeout=60,
        )
        if result.returncode == 0:
            base = os.path.splitext(os.path.basename(docx_path))[0]
            pdf_path = os.path.join(output_dir, base + ".pdf")
            if os.path.isfile(pdf_path):
                return pdf_path
    except Exception:
        pass
    return None


def export_pdf_docx2pdf(docx_path: str, pdf_path: str) -> Optional[str]:
    """Convert DOCX to PDF using docx2pdf (requires LO or MS Word)."""
    try:
        from docx2pdf import convert
        convert(docx_path, pdf_path)
        if os.path.isfile(pdf_path):
            return pdf_path
    except Exception:
        pass
    return None


def export_to_pdf(docx_path: str, timeout: int = 45) -> Optional[str]:
    """
    Try LibreOffice first, then docx2pdf — each in a thread with a timeout
    so a hanging Word/LO process can never crash the Flask request.
    Returns the PDF path or None if both fail.
    """
    import threading

    output_dir = os.path.dirname(docx_path)
    base = os.path.splitext(docx_path)[0]
    pdf_path = base + ".pdf"
    result_box: list[Optional[str]] = [None]

    def _try_lo():
        try:
            result_box[0] = export_pdf_libreoffice(docx_path, output_dir)
        except Exception:
            result_box[0] = None

    def _try_d2p():
        try:
            result_box[0] = export_pdf_docx2pdf(docx_path, pdf_path)
        except Exception:
            result_box[0] = None

    for fn in (_try_lo, _try_d2p):
        result_box[0] = None
        t = threading.Thread(target=fn, daemon=True)
        t.start()
        t.join(timeout=timeout)
        if result_box[0]:
            return result_box[0]

    return None


# ── High-level generate function ─────────────────────────────────────────────

def generate_outputs(
    template_path: str,
    sections: dict[str, str],
    company: str,
    role: str,
) -> dict[str, Optional[str]]:
    """
    Fill the template and produce both DOCX and PDF outputs.

    Returns:
        {
          "docx": "/path/to/output.docx",
          "pdf":  "/path/to/output.pdf" or None,
          "pdf_error": "reason" or None,
        }
    """
    # Filename format: AdityaMehta_CoverLetter_CompanyName
    safe_company = re.sub(r'[^\w\s-]', '', company).strip().replace(" ", "")
    filename = f"AdityaMehta_CoverLetter_{safe_company}"

    docx_path = os.path.join(config.OUTPUT_DIR, filename + ".docx")
    fill_template(template_path, sections, docx_path)

    pdf_path = export_to_pdf(docx_path)

    # Copy PDF to ~/Downloads so it's immediately accessible
    if pdf_path and os.path.isfile(pdf_path):
        downloads = os.path.expanduser("~/Downloads")
        downloads_copy = os.path.join(downloads, os.path.basename(pdf_path))
        try:
            shutil.copy2(pdf_path, downloads_copy)
        except Exception:
            pass

    pdf_error = None if pdf_path else (
        "LibreOffice not found. Install it free from https://www.libreoffice.org/ "
        "for automatic PDF export. Your DOCX is ready to download."
    )

    return {
        "docx": docx_path,
        "pdf": pdf_path,
        "pdf_error": pdf_error,
    }
