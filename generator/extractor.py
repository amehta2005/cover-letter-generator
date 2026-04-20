"""
extractor.py — Pull plain text from DOCX, PDF, and TXT files.

Handles the three formats present in this project:
  - .docx  (template, resume, and two example cover letters)
  - .pdf   (37 PDF cover letter examples)
  - .txt   (profile.txt)
"""

import os
import re
from docx import Document

# pdfplumber is optional at import time so the app still starts if it's missing
try:
    import pdfplumber
    _PDF_AVAILABLE = True
except ImportError:
    _PDF_AVAILABLE = False


# ── DOCX ──────────────────────────────────────────────────────────────────────

def extract_docx_text(path: str) -> str:
    """Return all paragraph text from a DOCX file, joined by newlines."""
    doc = Document(path)
    return "\n".join(p.text for p in doc.paragraphs if p.text.strip())


def extract_docx_paragraphs(path: str) -> list[str]:
    """Return non-empty paragraph strings from a DOCX file."""
    doc = Document(path)
    return [p.text.strip() for p in doc.paragraphs if p.text.strip()]


def extract_docx_sections(path: str) -> dict[str, str]:
    """
    Best-effort extraction of named sections from a resume DOCX.

    Looks for heading-style paragraphs (ALL CAPS or bold) and groups the
    body text that follows them.  Returns a dict like:
        {"EXPERIENCE": "...", "SKILLS": "...", "EDUCATION": "..."}
    """
    doc = Document(path)
    sections: dict[str, list[str]] = {}
    current = "HEADER"
    sections[current] = []

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue

        # Detect section headings: ALL-CAPS line OR bold paragraph style
        is_heading = (
            text.isupper()
            or (para.runs and all(r.bold for r in para.runs if r.text.strip()))
            or para.style.name.startswith("Heading")
        )

        if is_heading and len(text) < 60:
            current = text.upper()
            sections.setdefault(current, [])
        else:
            sections.setdefault(current, []).append(text)

    return {k: "\n".join(v) for k, v in sections.items() if v}


# ── PDF ───────────────────────────────────────────────────────────────────────

def extract_pdf_text(path: str) -> str:
    """Extract plain text from a PDF using pdfplumber."""
    if not _PDF_AVAILABLE:
        return ""
    try:
        with pdfplumber.open(path) as pdf:
            pages = [page.extract_text() or "" for page in pdf.pages]
        return "\n".join(pages)
    except Exception:
        return ""


# ── TXT ───────────────────────────────────────────────────────────────────────

def extract_txt(path: str) -> str:
    """Read a plain-text file."""
    with open(path, "r", encoding="utf-8") as f:
        return f.read()


# ── Examples folder ───────────────────────────────────────────────────────────

def load_examples(folder: str) -> list[str]:
    """
    Load all cover letter examples from a folder.
    Reads .docx first (better fidelity), then .pdf.
    Deduplicates by stem name so a file that exists as both
    .docx and .pdf is only loaded once (DOCX preferred).
    """
    if not os.path.isdir(folder):
        return []

    files = os.listdir(folder)
    seen_stems: set[str] = set()
    texts: list[str] = []

    # DOCX pass
    for f in sorted(files):
        if f.lower().endswith(".docx"):
            stem = os.path.splitext(f)[0]
            seen_stems.add(stem)
            text = extract_docx_text(os.path.join(folder, f))
            if text.strip():
                texts.append(text)

    # PDF pass — skip if DOCX already loaded for same stem
    for f in sorted(files):
        if f.lower().endswith(".pdf"):
            stem = os.path.splitext(f)[0]
            if stem in seen_stems:
                continue
            text = extract_pdf_text(os.path.join(folder, f))
            if text.strip():
                texts.append(text)
                seen_stems.add(stem)

    return texts


# ── Sentence mining from examples ─────────────────────────────────────────────

def mine_opener_sentences(example_texts: list[str]) -> list[str]:
    """
    Extract first sentences from each example cover letter.
    Used to anchor the tone of generated letters to the user's actual voice.
    """
    openers = []
    for text in example_texts:
        # First non-empty, non-date, non-address line that's a real sentence
        for line in text.splitlines():
            line = line.strip()
            if len(line) > 40 and not _looks_like_header(line):
                # Take only the first sentence
                sentence = re.split(r'(?<=[.!?])\s', line)[0]
                if len(sentence) > 30:
                    openers.append(sentence)
                break
    return openers


def mine_closing_sentences(example_texts: list[str]) -> list[str]:
    """Extract closing sentences from each example cover letter."""
    closings = []
    for text in example_texts:
        lines = [l.strip() for l in text.splitlines() if l.strip()]
        # Work backwards: find first line that looks like a real sentence
        for line in reversed(lines):
            if (
                len(line) > 30
                and not _looks_like_header(line)
                and line[0].isupper()
            ):
                closings.append(line)
                break
    return closings


def _looks_like_header(line: str) -> bool:
    """Heuristic: is this line a name/date/address header, not body text?"""
    short_words = len(line.split()) <= 4
    all_caps = line.isupper()
    looks_like_date = bool(re.search(r'\b(Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec|\d{4})\b', line))
    looks_like_contact = bool(re.search(r'[@|]\s*\d|linkedin|github|phone|email', line, re.I))
    return short_words or all_caps or looks_like_date or looks_like_contact
